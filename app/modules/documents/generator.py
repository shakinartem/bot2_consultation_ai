from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from sqlalchemy.ext.asyncio import AsyncSession

from app.config import get_settings
from app.modules.consultation.models import Consultation
from app.modules.crm.service import Company
from app.modules.files.storage import docs_storage_dir


BRAND_BLUE = RGBColor(6, 28, 65)
BRAND_BURGUNDY = RGBColor(118, 2, 41)
LIGHT_GRAY = RGBColor(231, 231, 231)


def _set_run_style(run, size: int = 11, bold: bool = False, color: RGBColor | None = None) -> None:
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def _add_title(document: Document, text: str, size: int = 16, color: RGBColor = BRAND_BLUE) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(10)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    _set_run_style(run, size=size, bold=True, color=color)


def _add_paragraph(document: Document, text: str, size: int = 11, bold: bool = False) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(5)
    run = paragraph.add_run(text)
    _set_run_style(run, size=size, bold=bold)


def _add_body(document: Document, text: str | None) -> None:
    if not text:
        text = "Нет данных. Нужно уточнить на консультации."
    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        if line.startswith(("- ", "* ")):
            paragraph = document.add_paragraph(style=None)
            paragraph.style = document.styles["List Bullet"]
            run = paragraph.add_run(line[2:].strip())
        else:
            paragraph = document.add_paragraph()
            run = paragraph.add_run(line.lstrip("#").strip())
        paragraph.paragraph_format.space_after = Pt(4)
        _set_run_style(run)


def _add_section(document: Document, title: str, text: str | None) -> None:
    _add_title(document, title, size=14)
    _add_body(document, text)


def _add_divider(document: Document) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run("•" * 42)
    _set_run_style(run, size=8, color=LIGHT_GRAY)


def _set_margins(document: Document) -> None:
    for section in document.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)


def _add_logo_if_available(document: Document) -> None:
    logo_path = get_settings().agency_logo_path
    if not logo_path.exists():
        return
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(logo_path), width=Inches(1.8))


def _add_cover(document: Document, company: Company) -> None:
    settings = get_settings()
    _set_margins(document)
    _add_logo_if_available(document)

    brand = document.add_paragraph()
    brand.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = brand.add_run(settings.agency_name)
    _set_run_style(run, size=28, bold=True, color=BRAND_BURGUNDY)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("AI-аудит и предварительное предложение")
    _set_run_style(run, size=20, bold=True, color=BRAND_BLUE)

    document.add_paragraph()
    target = document.add_paragraph()
    target.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = target.add_run(company.name)
    _set_run_style(run, size=17, bold=True)

    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(f"{company.city or 'Город не указан'} · {datetime.now().strftime('%d.%m.%Y')}")
    _set_run_style(run, size=12)

    lead = document.add_paragraph()
    lead.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = lead.add_run("Digital-система для стоматологических клиник")
    _set_run_style(run, size=11, color=BRAND_BLUE)
    document.add_page_break()


def _add_client_table(document: Document, company: Company) -> None:
    social = ", ".join(company.social_links) or "нет данных"
    rows = [
        ("Название", company.name),
        ("Юридическое название", company.legal_name or "нет данных"),
        ("Город", company.city or "нет данных"),
        ("Адрес", company.address or "нет данных"),
        ("Телефон", company.phone or "нет данных"),
        ("Сайт", company.website or "нет данных"),
        ("Карты", company.maps_url or "нет данных"),
        ("Соцсети", social),
        ("Рейтинг/отзывы", f"{company.rating or 'нет данных'} / {company.reviews_count or 'нет данных'}"),
        ("CRM заметки", company.crm_notes or "нет данных"),
    ]
    table = document.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for label, value in rows:
        cells = table.add_row().cells
        label_run = cells[0].paragraphs[0].add_run(label)
        value_run = cells[1].paragraphs[0].add_run(str(value))
        _set_run_style(label_run, bold=True, color=BRAND_BLUE)
        _set_run_style(value_run)


def _summary_text(consultation: Consultation) -> str:
    return "\n".join(
        item
        for item in (
            consultation.overall_conclusion,
            "Где клиника теряет заявки:",
            consultation.main_problems,
            "Что можно улучшить быстро:",
            consultation.quick_improvements,
        )
        if item
    )


def _agency_contacts_text() -> str:
    settings = get_settings()
    items = []
    if settings.agency_phone:
        items.append(f"Телефон: {settings.agency_phone}")
    if settings.agency_telegram:
        items.append(f"Telegram: {settings.agency_telegram}")
    if settings.agency_email:
        items.append(f"Email: {settings.agency_email}")
    if settings.agency_website:
        items.append(f"Сайт: {settings.agency_website}")
    if settings.agency_city:
        items.append(f"Город: {settings.agency_city}")
    if settings.agency_geo:
        items.append(f"География: {settings.agency_geo}")
    return "\n".join(items)


async def generate_consultation_docx(
    session: AsyncSession,
    consultation: Consultation,
    company: Company,
) -> Path:
    document = Document()
    _add_cover(document, company)

    _add_title(document, "Главный вывод", size=16, color=BRAND_BURGUNDY)
    _add_body(document, _summary_text(consultation))
    _add_divider(document)

    _add_title(document, "Данные клиента", size=16)
    _add_client_table(document, company)
    _add_divider(document)

    _add_section(document, "Контекст продаж", consultation.sales_context)
    _add_section(document, "Главный вывод", consultation.overall_conclusion)
    _add_section(document, "Где клиника теряет заявки", consultation.main_problems)
    _add_section(document, "Аудит сайта", consultation.website_audit)
    _add_section(document, "Аудит карт", consultation.maps_audit)
    _add_section(document, "Аудит соцсетей", consultation.social_audit)
    _add_section(document, "Репутация", consultation.reputation_audit)
    _add_section(document, "Основные проблемы", consultation.main_problems)
    _add_section(document, "Точки роста", consultation.growth_points)
    _add_section(document, "Roadmap 7 дней", consultation.roadmap_7_days)
    _add_section(document, "Roadmap 30 дней", consultation.roadmap_30_days)
    _add_section(document, "Roadmap 90 дней", consultation.roadmap_90_days)
    _add_section(document, "Что важно проговорить на консультации", consultation.consultation_talking_points)
    _add_section(document, "Предварительное предложение", consultation.proposal_text)
    _add_section(document, "Следующий шаг", consultation.next_step)

    document.add_section(WD_SECTION.CONTINUOUS)
    _add_divider(document)
    footer = document.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(f"{get_settings().agency_name} · Digital-система для стоматологических клиник")
    _set_run_style(run, size=10, bold=True, color=BRAND_BLUE)
    contacts = _agency_contacts_text()
    if contacts:
        _add_paragraph(document, contacts)

    output_path = docs_storage_dir() / f"consultation_company_{company.id}_{consultation.id}.docx"
    document.save(output_path)
    consultation.document_path = str(output_path)
    if consultation.proposal_text:
        consultation.proposal_document_path = str(output_path)
    await session.commit()
    await session.refresh(consultation)
    return output_path
